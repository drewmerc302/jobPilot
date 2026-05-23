"""Resume extraction: PDF/DOCX → structured profile dict.

Faithfully ports the resume-skills:extracting-resumes schema from
github.com/vibewith-brent/claude-resume-skills. Uses Anthropic's native
PDF document block for PDFs; python-docx text extraction for DOCX.
"""

import base64
import io
import logging

import anthropic
from pypdf import PdfReader

import jobpilot.llm as llm
from jobpilot.config import Config
from jobpilot.db import Database

logger = logging.getLogger(__name__)

# Schema ported from resume-extractor/references/resume_schema.yaml
EXTRACT_TOOL = {
    "name": "extract_resume",
    "description": "Extract structured profile data from a resume document",
    "input_schema": {
        "type": "object",
        "properties": {
            "contact": {
                "type": "object",
                "properties": {
                    "name": {"type": "string"},
                    "title": {
                        "type": "string",
                        "description": "Professional title / current role",
                    },
                    "email": {"type": "string"},
                    "phone": {"type": "string"},
                    "location": {"type": "string"},
                    "linkedin": {"type": "string"},
                    "github": {"type": "string"},
                    "website": {"type": "string"},
                },
                "required": ["name", "email"],
            },
            "summary": {
                "type": "string",
                "description": "Professional summary paragraph, 2-5 sentences",
            },
            "experience": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "company": {"type": "string"},
                        "location": {"type": "string"},
                        "positions": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "title": {"type": "string"},
                                    "dates": {"type": "string"},
                                    "achievements": {
                                        "type": "array",
                                        "items": {"type": "string"},
                                    },
                                },
                                "required": ["title", "dates", "achievements"],
                            },
                        },
                    },
                    "required": ["company", "positions"],
                },
            },
            "skills": {
                "type": "object",
                "description": "Skills organised by category name → list of skill strings",
                "additionalProperties": {
                    "type": "array",
                    "items": {"type": "string"},
                },
            },
            "education": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "institution": {"type": "string"},
                        "degree": {"type": "string"},
                        "graduation_year": {"type": "string"},
                        "location": {"type": "string"},
                        "gpa": {"type": "string"},
                        "honors": {"type": "string"},
                        "minor": {"type": "string"},
                    },
                    "required": ["institution", "degree"],
                },
            },
            "certifications": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "name": {"type": "string"},
                        "issuer": {"type": "string"},
                        "date": {"type": "string"},
                    },
                    "required": ["name", "issuer"],
                },
            },
            "low_confidence_fields": {
                "type": "array",
                "items": {"type": "string"},
                "description": (
                    "Field paths where extraction confidence is low due to ambiguity, "
                    "illegibility, or inference (e.g. 'contact.email', 'experience[0].positions[0].dates')"
                ),
            },
        },
        "required": [
            "contact",
            "experience",
            "skills",
            "education",
            "low_confidence_fields",
        ],
    },
}

_EXTRACTION_PROMPT = """Extract all resume content into the structured format using the extract_resume tool.

Follow these rules:
- contact.name is required; mark in low_confidence_fields if uncertain
- Never use placeholders like "<UNKNOWN>" or "N/A" — make your best guess from context, or leave the field as an empty string. Add the field path to low_confidence_fields instead.
- contact.title: the person's professional title or most recent role descriptor
- experience: list reverse-chronologically (most recent first)
- Each company may have multiple positions for promotions/role changes
- achievements: start each bullet with an action verb; preserve metrics exactly
- skills: group by category (e.g. "Programming Languages", "Frameworks", "Tools")
- dates: normalise to "Mon YYYY - Mon YYYY" or "Mon YYYY - Present"
- education: reverse-chronological
- linkedin, github, website: extract full URLs (e.g. https://linkedin.com/in/...), not display text
- low_confidence_fields: list any field you are uncertain about
- Omit optional sections (certifications, etc.) if not present in the resume"""


def _extract_pdf_links(file_bytes: bytes) -> list[str]:
    """Pull all hyperlink URIs embedded in a PDF's annotations."""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        urls: list[str] = []
        for page in reader.pages:
            for annot in page.get("/Annots") or []:
                obj = annot.get_object()
                if obj.get("/Subtype") == "/Link":
                    uri = (obj.get("/A") or {}).get("/URI")
                    if uri and uri not in urls:
                        urls.append(str(uri))
        return urls
    except Exception:
        return []


def _extract_text_from_docx(file_bytes: bytes) -> str:
    import docx  # python-docx

    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also pull text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def extract_resume(
    file_bytes: bytes,
    filename: str,
    client: anthropic.Anthropic,
    config: Config,
    db: Database,
) -> dict:
    """Extract structured profile from PDF or DOCX bytes.

    Returns a profile dict matching the resume schema.
    Raises ValueError if LLM returns no tool_use block.
    """
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        b64 = base64.standard_b64encode(file_bytes).decode()
        links = _extract_pdf_links(file_bytes)
        link_hint = ""
        if links:
            link_hint = (
                "\n\nThe PDF contains these embedded hyperlinks — use them for "
                "linkedin, github, and website fields:\n"
                + "\n".join(f"- {u}" for u in links)
            )
        content = [
            {
                "type": "document",
                "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": b64,
                },
            },
            {"type": "text", "text": _EXTRACTION_PROMPT + link_hint},
        ]
    elif filename_lower.endswith(".docx"):
        text = _extract_text_from_docx(file_bytes)
        content = [
            {"type": "text", "text": f"RESUME TEXT:\n\n{text}\n\n{_EXTRACTION_PROMPT}"},
        ]
    else:
        raise ValueError(f"Unsupported file type: {filename!r}. Use PDF or DOCX.")

    response = llm.call(
        client,
        db,
        "extract_resume",
        model=config.llm_extract_model,
        max_tokens=4096,
        tools=[EXTRACT_TOOL],
        tool_choice={"type": "tool", "name": "extract_resume"},
        messages=[{"role": "user", "content": content}],
    )

    for block in response.content:
        if hasattr(block, "type") and block.type == "tool_use":
            profile = block.input
            # Flatten: hoist contact fields to top level for easier template access
            contact = profile.pop("contact", {})
            profile["name"] = contact.get("name", "")
            profile["email"] = contact.get("email", "")
            profile["phone"] = contact.get("phone", "")
            profile["location"] = contact.get("location", "")
            profile["linkedin"] = contact.get("linkedin", "")
            profile["github"] = contact.get("github", "")
            profile["website"] = contact.get("website", "")
            profile["title"] = contact.get("title", "")
            # B7.3: defensive defaults so downstream templates and edits
            # do not crash on a partially-populated extraction.
            profile.setdefault("experience", [])
            profile.setdefault("skills", {})
            profile.setdefault("education", [])
            profile.setdefault("low_confidence_fields", [])
            return profile

    raise ValueError("LLM did not return a structured resume extraction")
